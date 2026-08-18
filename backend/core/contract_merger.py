"""
contract_merger.py — 合并已采纳的修订，生成修订后合同。

纯函数模块：读取 clauses + 已采纳(accepted)的 revisions，用 after_text 替换
对应条款，输出合并后的合同全文（纯文本）或 .docx 文件。

设计原则：
- 与审查 agent 通过数据库解耦。审查 agent 产出 revision（含 before/after_text）
  落库，本模块只消费 status='accepted' 的修订做**确定性文本替换**，不涉及 LLM。
- 关联靠 clause_id（修订 bug 后为真实 hash id），不再用模糊匹配。
- 未采纳/驳回/需律师的修订一律不动原条款，保持原文。
- 导出时过滤 LLM 误填的辅助标签（合同标题/当事人信息/签署落款等），
  只输出真实条款标题与正文。
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)

# ── 标签过滤 ────────────────────────────────────────────────

# LLM 抽取时给「标题/声明/落款」等非实质条款填的辅助标签，不应出现在合同正文
_LABEL_TITLES = {
    "合同标题", "合同编号", "当事人信息", "当事人", "前言", "前言声明",
    "鉴于", "签署栏", "签署落款", "落款", "标题", "声明", "定义",
}

# 真实条款标题特征：含「第X条/章/节」
_CLAUSE_HEADING_RE = re.compile(r"第[一二三四五六七八九十百千零]+[条章节]")


def _is_label_title(title: str) -> bool:
    """判断 title 是否为机器标签（非真实条款标题）。"""
    t = (title or "").strip()
    if not t:
        return False
    if _CLAUSE_HEADING_RE.search(t):
        return False  # 真实条款标题
    if t in _LABEL_TITLES:
        return True
    return len(t) <= 4  # 兜底：短且无条款号 → 视为标签候选


def _is_contract_title_clause(c: dict[str, Any]) -> bool:
    """判断 clause 是否为「合同标题」类（content 是合同主标题）。"""
    title = (c.get("title", "") or "").strip()
    content = (c.get("content", "") or "").strip()
    if title == "合同标题" and content:
        return True
    # 兜底：seq_no==1 且 content 是短标题、title 是标签或空 → 视为标题类
    if c.get("seq_no") == 1 and content and len(content) <= 20 and not _CLAUSE_HEADING_RE.search(content):
        if not title or _is_label_title(title):
            return True
    return False


def _extract_contract_title(clauses: list[dict[str, Any]]) -> str | None:
    """从 clauses 提取真实合同标题（「合同标题」类 clause 的 content）。"""
    for c in sorted(clauses, key=lambda x: x.get("seq_no", 0)):
        if _is_contract_title_clause(c):
            return (c.get("content", "") or "").strip()
    return None


# ── Merge ───────────────────────────────────────────────────

def merge_clauses(
    clauses: list[dict[str, Any]],
    accepted_revisions: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """按 seq_no 遍历条款，accepted 的用 after_text 替换 content。

    Returns:
        list of {seq_no, title, content, changed}
    """
    accepted_map: dict[str, str] = {
        r.get("clause_id", ""): r.get("after_text", "")
        for r in accepted_revisions
        if r.get("clause_id")
    }

    merged: list[dict[str, Any]] = []
    for c in sorted(clauses, key=lambda x: x.get("seq_no", 0)):
        cid = c.get("clause_id", "")
        new_content = accepted_map.get(cid, c.get("content", ""))
        merged.append({
            "seq_no": c.get("seq_no", 0),
            "title": c.get("title", ""),
            "content": new_content,
            "changed": cid in accepted_map,
        })
    return merged


# ── Text output ─────────────────────────────────────────────

def build_contract_text(
    merged_clauses: list[dict[str, Any]],
    contract_title: str | None = None,
) -> str:
    """拼接合并后的合同全文（纯文本，条款间空行分隔，过滤标签）。"""
    real_title = _extract_contract_title(merged_clauses) or contract_title

    lines: list[str] = []
    if real_title:
        lines.append(real_title)
        lines.append("")
    for c in merged_clauses:
        if _is_contract_title_clause(c):
            continue  # 标题类已提取为主标题，跳过
        title = (c.get("title", "") or "").strip()
        if title and not _is_label_title(title):
            lines.append(title)
        if c.get("content"):
            lines.append(c["content"])
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


# ── Docx output ─────────────────────────────────────────────

def _set_font(run, name: str, size=None) -> None:
    """设置 run 的字体（含中文字体 eastAsia，否则中文回退默认字体）。"""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def _add_page_number(paragraph) -> None:
    """向段落追加「PAGE」页码域。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    _set_font(run, "宋体", 9)


def build_contract_docx(
    merged_clauses: list[dict[str, Any]],
    contract_title: str | None = None,
) -> bytes:
    """生成修订后合同的 .docx 文件字节流（中文合同排版规范）。

    排版：主标题黑体 22pt 居中；条款标题黑体 12pt 加粗；正文仿宋 12pt、
    1.5 倍行距、首行缩进 2 字符；页眉合同名、页脚页码。
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    # 真实标题：优先从 clauses 的「合同标题」提取，否则用传入的文件名标题
    real_title = _extract_contract_title(merged_clauses) or contract_title

    # 页眉：合同名（宋体 9pt 居中）
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(hp.add_run(real_title or ""), "宋体", 9)

    # 页脚：页码（居中）
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(fp)

    # 主标题（黑体 22pt 居中）
    if real_title:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(real_title)
        run.bold = True
        _set_font(run, "黑体", 22)

    for c in merged_clauses:
        if _is_contract_title_clause(c):
            continue  # 标题类已提取为主标题，跳过

        title = (c.get("title", "") or "").strip()
        content = c.get("content", "") or ""

        # 条款标题（黑体 12pt 加粗，真实条款才输出）
        if title and not _is_label_title(title):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(title)
            run.bold = True
            _set_font(run, "黑体", 12)

        # 正文（仿宋 12pt，1.5 倍行距，首行缩进 2 字符）
        for para in content.split("\n"):
            if not para.strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Pt(24)  # 12pt × 2 字符
            _set_font(p.add_run(para.strip()), "仿宋", 12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
