"""
Contract review API routes — /contract/*
Upload, run, stream, report, revision accept, lawyer confirm.
"""
from __future__ import annotations

import asyncio
import json
import logging
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import asyncpg
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from urllib.parse import quote

from backend.agents.contract_review.graph import run_contract_review, run_contract_review_direct
from backend.config import get_settings
from backend.core.contract_merger import merge_clauses, build_contract_text, build_contract_docx
from backend.dependencies import get_current_user, decode_token
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials

# ── Review cards store ──────────────────────────────────────

def _save_review_cards(contract_id: int, cards: list[dict[str, Any]]) -> None:
    pass  # now persisted to DB via _persist_review_cards


async def _load_review_cards_from_db(contract_id: int) -> list[dict[str, Any]]:
    """Load review cards from contract_review_cards table (survives restarts)."""
    db = await _get_db()
    try:
        rows = await db.fetch(
            "SELECT * FROM contract_review_cards WHERE review_id = $1",
            contract_id,
        )
        return [dict(r) for r in rows]
    finally:
        await db.close()


logger = logging.getLogger(__name__)

router = APIRouter(prefix="/contract", tags=["contract"])

UPLOAD_DIR = Path("uploads/contracts")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# ── Database helper ─────────────────────────────────────────

DATABASE_URL = get_settings().database_url


async def _get_db() -> asyncpg.Connection:
    return await asyncpg.connect(DATABASE_URL)


# ── Pydantic schemas ────────────────────────────────────────

class ContractUploadResponse(BaseModel):
    contract_id: int
    filename: str


class RunResponse(BaseModel):
    run_id: str
    contract_id: int
    status: str


class RevisionAcceptRequest(BaseModel):
    status: str  # accepted / rejected / needs_lawyer
    idempotent_key: str


class LawyerConfirmRequest(BaseModel):
    confirmed: bool


class HumanDecisionRequest(BaseModel):
    """Request body for submitting human review decisions."""
    decisions: list[dict[str, Any]]  # [{clause_id, action, modified_level, modified_score, comment, skip_revision}, ...]


# ── SSE helpers ─────────────────────────────────────────────

async def _sse_event(event: str, data: Any) -> str:
    """Format an SSE event string."""
    payload = json.dumps(data, ensure_ascii=False)
    return f"event: {event}\ndata: {payload}\n\n"


async def _resolve_current_user(
    contract_id: int | None = None,
    credentials: HTTPAuthorizationCredentials | None = None,
    token: str | None = None,
) -> dict:
    """Resolve current user from Authorization header OR query param ?token=.

    The SSE stream endpoint can't easily send Authorization headers from the
    browser (EventSource limitation), so it falls back to ?token=. For all
    other endpoints we prefer the Authorization header.
    """
    user = None
    if credentials:
        try:
            user = await get_current_user_from_credentials(credentials)
        except HTTPException:
            pass
    if user is None and token:
        try:
            from backend.dependencies import decode_token
            payload = decode_token(token)
            user = {
                "user_id": int(payload.get("sub", 0)),
                "username": payload.get("username", ""),
                "role": payload.get("role", "user"),
            }
        except HTTPException:
            pass
    if user is None:
        raise HTTPException(401, "Authentication required")
    return user


async def get_current_user_from_credentials(credentials: HTTPAuthorizationCredentials) -> dict:
    """Decode JWT from HTTPBearer credentials only (no fallback)."""
    payload = decode_token(credentials.credentials)
    return {
        "user_id": int(payload.get("sub", 0)),
        "username": payload.get("username", ""),
        "role": payload.get("role", "user"),
    }


# ── Helpers ─────────────────────────────────────────────────

async def _check_ownership(
    db: asyncpg.Connection, contract_id: int, user_id: int, is_admin: bool
) -> None:
    """Verify the requesting user owns the contract; raises HTTPException if not."""
    review = await db.fetchrow(
        "SELECT user_id FROM contract_reviews WHERE id = $1", contract_id
    )
    if not review:
        raise HTTPException(404, "Contract not found")
    if review["user_id"] != user_id and not is_admin:
        raise HTTPException(403, "无权操作此审查")


# ── Routes ──────────────────────────────────────────────────

@router.post("/upload", response_model=ContractUploadResponse)
async def upload_contract(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
):
    """Upload a PDF/DOCX contract file. Returns contract_id."""
    if not file.filename:
        raise HTTPException(400, "Filename required")

    # Validate extension
    ext = Path(file.filename).suffix.lower()
    if ext not in (".pdf", ".docx", ".txt"):
        raise HTTPException(400, f"Unsupported file type: {ext}. Only PDF, DOCX, TXT allowed.")

    # Save file
    file_id = uuid.uuid4().hex[:12]
    saved_name = f"{file_id}_{file.filename}"
    saved_path = UPLOAD_DIR / saved_name

    content = await file.read()
    saved_path.write_bytes(content)

    # Extract text
    text = ""
    if ext == ".txt":
        text = content.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(str(saved_path)) as pdf:
                parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        parts.append(page_text)
                text = "\n\n".join(parts)
        except ImportError:
            text = "[PDF parsing requires pdfplumber]"
    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(saved_path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(parts)
        except ImportError:
            text = "[DOCX parsing requires python-docx]"

    if not text.strip():
        raise HTTPException(400, "Could not extract text from the uploaded file")

    user_id = current_user["user_id"]

    # Insert into database
    db = await _get_db()
    try:
        row = await db.fetchrow(
            """
            INSERT INTO contract_reviews (user_id, filename, original_filename, contract_type, status)
            VALUES ($1, $2, $3, '', 'pending')
            RETURNING id
            """,
            user_id, saved_name, file.filename,
        )
        contract_id = row["id"]
        logger.info(f"Contract uploaded: id={contract_id}, file={file.filename}, user_id={user_id}")
    finally:
        await db.close()

    return ContractUploadResponse(contract_id=contract_id, filename=file.filename)


@router.post("/run/{contract_id}", response_model=RunResponse)
async def run_contract_review_endpoint(
    contract_id: int,
    current_user: dict = Depends(get_current_user),
):
    """Trigger async contract review. Returns run_id for tracking."""
    user_id = current_user["user_id"]
    db = await _get_db()
    try:
        review = await db.fetchrow(
            "SELECT id, filename, user_id FROM contract_reviews WHERE id = $1", contract_id
        )
        if not review:
            raise HTTPException(404, "Contract not found")
        if review["user_id"] != user_id and current_user.get("role") != "admin":
            raise HTTPException(403, "无权操作此审查")

        # Read the uploaded file
        saved_path = UPLOAD_DIR / review["filename"]
        if not saved_path.exists():
            raise HTTPException(404, "Uploaded file not found on disk")

        # Parse text (same logic as upload)
        ext = Path(review["filename"]).suffix.lower()
        text = ""
        if ext == ".txt":
            text = saved_path.read_text(encoding="utf-8", errors="replace")
        elif ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(str(saved_path)) as pdf:
                    parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            parts.append(page_text)
                    text = "\n\n".join(parts)
            except ImportError:
                text = "[PDF parsing requires pdfplumber]"
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(str(saved_path))
                parts = [p.text for p in doc.paragraphs if p.text.strip()]
                text = "\n\n".join(parts)
            except ImportError:
                text = "[DOCX parsing requires python-docx]"

        if not text.strip():
            raise HTTPException(400, "Could not extract text from file")

        # Update status
        await db.execute(
            "UPDATE contract_reviews SET status = 'parsing', updated_at = NOW() WHERE id = $1",
            contract_id,
        )
    finally:
        await db.close()

    run_id = uuid.uuid4().hex[:8]

    status_map = {
        "parse": "parsing",
        "rule_check": "parsing",  # rule_check shares parsing stage in SSE progress
        "review": "reviewing",
        "retrieve": "retrieving",
        "revise": "revising",
    }

    async def _progress_callback(node_name: str, event: str, data: dict[str, Any]):
        """Update DB status on each node start for SSE streaming."""
        if event == "paused":
            # Human-in-the-Loop: pipeline paused for review
            db2 = await _get_db()
            try:
                await db2.execute(
                    "UPDATE contract_reviews SET status = 'paused_waiting', updated_at = NOW() WHERE id = $1",
                    contract_id,
                )
            finally:
                await db2.close()
        elif event == "started" and node_name in status_map:
            new_status = status_map[node_name]
            db2 = await _get_db()
            try:
                await db2.execute(
                    "UPDATE contract_reviews SET status = $2, updated_at = NOW() WHERE id = $1",
                    contract_id, new_status,
                )
            finally:
                await db2.close()

    async def _run_pipeline():
        review_cards = []
        try:
            result = await run_contract_review_direct(
                contract_id=contract_id,
                user_id=user_id,
                text=text,
                progress_callback=_progress_callback,
            )
            review_cards = result.get("review_cards", [])
            if review_cards:
                _save_review_cards(contract_id, review_cards)
            db2 = await _get_db()
            try:
                await db2.execute(
                    "UPDATE contract_reviews SET status = 'completed', updated_at = NOW() WHERE id = $1",
                    contract_id,
                )
            finally:
                await db2.close()
        except Exception as e:
            logger.exception(f"Pipeline failed for contract {contract_id}")
            db2 = await _get_db()
            try:
                await db2.execute(
                    "UPDATE contract_reviews SET status = 'failed', error_message = $2, updated_at = NOW() WHERE id = $1",
                    contract_id, str(e)[:500],
                )
            finally:
                await db2.close()

    asyncio.create_task(_run_pipeline())

    return RunResponse(run_id=run_id, contract_id=contract_id, status="started")


@router.get("/run/{contract_id}/stream")
async def stream_progress(
    contract_id: int,
    token: str = Query(""),
    credentials: HTTPAuthorizationCredentials | None = Depends(HTTPBearer(auto_error=False)),
):
    """SSE endpoint for real-time progress streaming.

    Accepts authentication via either the Authorization: Bearer <token>
    header OR a ?token=<jwt> query parameter (the latter is needed because
    browser EventSource cannot set custom headers).
    """
    current_user = await _resolve_current_user(
        contract_id=contract_id, credentials=credentials, token=token or None
    )
    user_id = current_user["user_id"]
    is_admin = current_user.get("role") == "admin"

    # Verify ownership
    db_check = await _get_db()
    try:
        review = await db_check.fetchrow(
            "SELECT id, user_id FROM contract_reviews WHERE id = $1", contract_id
        )
        if not review:
            raise HTTPException(404, "Contract not found")
        if review["user_id"] != user_id and not is_admin:
            raise HTTPException(403, "无权监控此审查")
    finally:
        await db_check.close()

    async def _event_stream():
        db = await _get_db()
        try:
            last_status = "pending"
            while True:
                row = await db.fetchrow(
                    "SELECT status, error_message FROM contract_reviews WHERE id = $1",
                    contract_id,
                )
                if not row:
                    yield await _sse_event("error", {"message": "Contract not found"})
                    return

                current_status = row["status"]
                if current_status != last_status:
                    yield await _sse_event("progress", {
                        "status": current_status,
                        "contract_id": contract_id,
                    })
                    last_status = current_status

                if current_status in ("completed", "failed", "paused_waiting"):
                    event_type = "paused" if current_status == "paused_waiting" else "complete"
                    yield await _sse_event(event_type, {
                        "status": current_status,
                        "contract_id": contract_id,
                        "error": row["error_message"],
                    })
                    if current_status in ("completed", "failed"):
                        return

                await asyncio.sleep(1)
        finally:
            await db.close()

    return StreamingResponse(
        _event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.get("/report/{contract_id}")
async def get_report(
    contract_id: int,
    current_user: dict = Depends(get_current_user),
):
    """Get complete review report including disclaimer."""
    db = await _get_db()
    user_id = current_user["user_id"]
    is_admin = current_user.get("role") == "admin"
    try:
        review = await db.fetchrow(
            "SELECT * FROM contract_reviews WHERE id = $1", contract_id
        )
        if not review:
            raise HTTPException(404, "Contract not found")
        if review["user_id"] != user_id and not is_admin:
            raise HTTPException(403, "无权查看此审查报告")

        clauses = await db.fetch(
            "SELECT * FROM contract_clauses WHERE review_id = $1 ORDER BY seq_no",
            contract_id,
        )
        review_cards = await db.fetch(
            "SELECT * FROM contract_review_cards WHERE review_id = $1",
            contract_id,
        )
        evidence = await db.fetch(
            "SELECT * FROM contract_evidence WHERE review_id = $1",
            contract_id,
        )
        revisions = await db.fetch(
            "SELECT * FROM revision_accepts WHERE review_id = $1",
            contract_id,
        )
        rule_findings = await db.fetch(
            "SELECT * FROM rule_findings WHERE review_id = $1",
            contract_id,
        )
        human_decisions = await db.fetch(
            "SELECT * FROM human_review_decisions WHERE review_id = $1",
            contract_id,
        )

        return {
            "contract_id": contract_id,
            "status": review["status"],
            "contract_type": review["contract_type"],
            "filename": review["original_filename"],
            "disclaimer_accepted": review["disclaimer_accepted"],
            "lawyer_confirmed_at": str(review["lawyer_confirmed_at"]) if review["lawyer_confirmed_at"] else None,
            "needs_human_review": review.get("needs_human_review", False),
            "human_review_status": review.get("human_review_status", "skipped"),
            "clauses": [dict(c) for c in clauses],
            "review_cards": [dict(r) for r in review_cards],
            "evidence": [dict(e) for e in evidence],
            "revisions": [dict(r) for r in revisions],
            "rule_findings": [dict(rf) for rf in rule_findings],
            "human_review_decisions": [dict(hd) for hd in human_decisions],
            "disclaimer": (
                "⚠ 免责声明：本报告由AI自动生成，仅供参考，不构成法律意见。"
                "所有审查结果需经执业律师审核确认后方可作为决策依据。"
                "使用本系统即表示您已理解并同意上述声明。"
            ),
            "created_at": str(review["created_at"]),
            "updated_at": str(review["updated_at"]) if review["updated_at"] else None,
        }
    finally:
        await db.close()


@router.get("/reviews")
async def list_reviews(current_user: dict = Depends(get_current_user)):
    """List recent contract reviews with risk-level counts (审查历史)."""
    db = await _get_db()
    user_id = current_user["user_id"]
    is_admin = current_user.get("role") == "admin"
    logger.warning(f"DEBUG list_reviews: user_id={user_id}, is_admin={is_admin}")
    try:
        if is_admin:
            rows = await db.fetch(
                "SELECT id, original_filename, contract_type, status, created_at, updated_at, user_id "
                "FROM contract_reviews ORDER BY id DESC LIMIT 50"
            )
        else:
            rows = await db.fetch(
                "SELECT id, original_filename, contract_type, status, created_at, updated_at, user_id "
                "FROM contract_reviews WHERE user_id = $1 ORDER BY id DESC LIMIT 50",
                user_id,
            )
        result = []
        for r in rows:
            counts = await db.fetch(
                "SELECT level, COUNT(*) AS n FROM contract_review_cards "
                "WHERE review_id = $1 GROUP BY level",
                r["id"],
            )
            lv = {c["level"]: c["n"] for c in counts}
            result.append({
                "id": r["id"],
                "filename": r["original_filename"],
                "contract_type": r["contract_type"],
                "status": r["status"],
                "created_at": str(r["created_at"]),
                "updated_at": str(r["updated_at"]) if r["updated_at"] else None,
                "high_risk": lv.get("高", 0),
                "medium_risk": lv.get("中", 0),
                "low_risk": lv.get("低", 0),
            })
        return {"reviews": result}
    finally:
        await db.close()


# Statuses that mean the pipeline is still running — deleting mid-run would
# leave the background task writing into a vanished review row.
_ACTIVE_STATUSES = ("pending", "parsing", "reviewing", "retrieving", "revising", "paused_waiting")


@router.delete("/review/{contract_id}")
async def delete_review(
    contract_id: int,
    current_user: dict = Depends(get_current_user),
):
    """Delete a historical review and all of its dependent data.

    All child tables (contract_clauses, contract_review_cards, contract_evidence,
    revision_accepts, contract_qa_sessions → contract_qa_messages) are removed in
    a single transaction. The uploaded file itself is kept on disk.
    """
    user_id = current_user["user_id"]
    is_admin = current_user.get("role") == "admin"
    db = await _get_db()
    try:
        review = await db.fetchrow(
            "SELECT id, status, user_id FROM contract_reviews WHERE id = $1", contract_id
        )
        if not review:
            raise HTTPException(404, "Contract not found")
        if review["user_id"] != user_id and not is_admin:
            raise HTTPException(403, "无权删除此审查")
        if review["status"] in _ACTIVE_STATUSES:
            raise HTTPException(409, "该审查正在进行中，请等待完成后再删除")

        async with db.transaction():
            await db.execute(
                "DELETE FROM contract_qa_messages WHERE session_id IN "
                "(SELECT id FROM contract_qa_sessions WHERE contract_id = $1)",
                contract_id,
            )
            await db.execute(
                "DELETE FROM contract_qa_sessions WHERE contract_id = $1", contract_id
            )
            await db.execute(
                "DELETE FROM contract_evidence WHERE review_id = $1", contract_id
            )
            await db.execute(
                "DELETE FROM revision_accepts WHERE review_id = $1", contract_id
            )
            await db.execute(
                "DELETE FROM contract_review_cards WHERE review_id = $1", contract_id
            )
            await db.execute(
                "DELETE FROM contract_clauses WHERE review_id = $1", contract_id
            )
            await db.execute(
                "DELETE FROM contract_reviews WHERE id = $1", contract_id
            )

        logger.info(f"Review deleted: id={contract_id}")
        return {"status": "deleted", "id": contract_id}
    finally:
        await db.close()


@router.post("/revision/{revision_id}/accept")
async def accept_revision(
    revision_id: int,
    req: RevisionAcceptRequest,
    current_user: dict = Depends(get_current_user),
):
    """Accept/reject a revision suggestion with idempotent key."""
    if req.status not in ("accepted", "rejected", "needs_lawyer"):
        raise HTTPException(400, "Invalid status. Must be accepted, rejected, or needs_lawyer.")

    user_id = current_user["user_id"]
    is_admin = current_user.get("role") == "admin"

    db = await _get_db()
    try:
        # Verify revision exists and user owns the parent review
        rev = await db.fetchrow(
            "SELECT rv.review_id, cr.user_id FROM revision_accepts rv "
            "JOIN contract_reviews cr ON cr.id = rv.review_id "
            "WHERE rv.id = $1", revision_id
        )
        if not rev:
            raise HTTPException(404, "Revision not found")
        if rev["user_id"] != user_id and not is_admin:
            raise HTTPException(403, "无权操作此修订")
        existing = await db.fetchrow(
            "SELECT id FROM idempotent_ops WHERE idempotent_key = $1",
            req.idempotent_key,
        )
        if existing:
            logger.info(f"Idempotent key already processed: {req.idempotent_key}")
            return {"status": "already_processed", "revision_id": revision_id}

        # Record idempotent op
        await db.execute(
            "INSERT INTO idempotent_ops (idempotent_key, operation_type, status) VALUES ($1, 'revision_accept', 'completed')",
            req.idempotent_key,
        )

        # Update revision
        await db.execute(
            "UPDATE revision_accepts SET status = $1, updated_at = NOW() WHERE id = $2",
            req.status, revision_id,
        )

        return {"status": "ok", "revision_id": revision_id, "new_status": req.status}
    finally:
        await db.close()


@router.post("/revision/{revision_id}/lawyer-confirm")
async def lawyer_confirm_revision(
    revision_id: int,
    req: LawyerConfirmRequest,
    current_user: dict = Depends(get_current_user),
):
    """Lawyer confirms a revision (professional review step)."""
    user_id = current_user["user_id"]
    is_admin = current_user.get("role") == "admin"
    db = await _get_db()
    try:
        rev = await db.fetchrow(
            "SELECT rv.review_id, cr.user_id FROM revision_accepts rv "
            "JOIN contract_reviews cr ON cr.id = rv.review_id "
            "WHERE rv.id = $1", revision_id
        )
        if not rev:
            raise HTTPException(404, "Revision not found")
        if rev["user_id"] != user_id and not is_admin:
            raise HTTPException(403, "无权操作此修订")

        if req.confirmed:
            await db.execute(
                "UPDATE contract_reviews SET lawyer_confirmed_at = NOW() WHERE id = $1",
                rev["review_id"],
            )

        return {"status": "ok", "revision_id": revision_id, "confirmed": req.confirmed}
    finally:
        await db.close()


# ── Final contract export ───────────────────────────────────

async def _load_contract_for_merge(contract_id: int, user_id: int, is_admin: bool):
    """加载审查 + 条款 + 修订，校验所有权。返回 (review, clauses, revisions)。"""
    db = await _get_db()
    review = await db.fetchrow(
        "SELECT * FROM contract_reviews WHERE id = $1", contract_id
    )
    if not review:
        await db.close()
        raise HTTPException(404, "Contract not found")
    if review["user_id"] != user_id and not is_admin:
        await db.close()
        raise HTTPException(403, "无权查看此审查报告")
    clauses = await db.fetch(
        "SELECT * FROM contract_clauses WHERE review_id = $1 ORDER BY seq_no",
        contract_id,
    )
    revisions = await db.fetch(
        "SELECT * FROM revision_accepts WHERE review_id = $1", contract_id
    )
    return review, clauses, revisions, db


@router.get("/{contract_id}/final-contract")
async def get_final_contract(
    contract_id: int,
    current_user: dict = Depends(get_current_user),
):
    """预览修订后合同全文与决策统计（供前端展示/校验）。"""
    user_id = current_user["user_id"]
    is_admin = current_user.get("role") == "admin"
    review, clauses, revisions, db = await _load_contract_for_merge(
        contract_id, user_id, is_admin
    )
    try:
        accepted = [dict(r) for r in revisions if r["status"] == "accepted"]
        pending_count = sum(1 for r in revisions if r["status"] == "pending")

        merged = merge_clauses([dict(c) for c in clauses], accepted)
        title = (review["original_filename"] or "合同").rsplit(".", 1)[0]
        text = build_contract_text(merged, title)

        return {
            "contract_id": contract_id,
            "title": title,
            "text": text,
            "pending_count": pending_count,
            "accepted_count": len(accepted),
            "total_revisions": len(revisions),
            "changed_count": sum(1 for m in merged if m["changed"]),
            "ready": pending_count == 0,
        }
    finally:
        await db.close()


@router.get("/{contract_id}/final-contract/download")
async def download_final_contract(
    contract_id: int,
    current_user: dict = Depends(get_current_user),
):
    """下载修订后合同的 .docx 文件；有未决策修订时拒绝并提示。"""
    user_id = current_user["user_id"]
    is_admin = current_user.get("role") == "admin"
    review, clauses, revisions, db = await _load_contract_for_merge(
        contract_id, user_id, is_admin
    )
    try:
        pending_count = sum(1 for r in revisions if r["status"] == "pending")
        if pending_count > 0:
            raise HTTPException(
                409, f"还有 {pending_count} 条修订未决策，请先全部决策后再导出"
            )

        accepted = [dict(r) for r in revisions if r["status"] == "accepted"]
        merged = merge_clauses([dict(c) for c in clauses], accepted)
        title = (review["original_filename"] or "合同").rsplit(".", 1)[0]

        docx_bytes = build_contract_docx(merged, title)
        filename = f"{title}_修订后.docx"
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
        )
    finally:
        await db.close()


# ── Human-in-the-Loop endpoint ──────────────────────────────

@router.post("/{contract_id}/human-decision")
async def submit_human_decision(
    contract_id: int,
    req: HumanDecisionRequest,
    current_user: dict = Depends(get_current_user),
):
    """Submit human review decisions for high-risk items and resume the pipeline.

    When the pipeline is paused at the human_gate, this endpoint saves the
    decisions to the DB and attempts to resume execution.

    In the current architecture (direct execution path), this records the
    decisions and marks the review as ready to continue. The actual pipeline
    resume is triggered by re-running with HITL decisions pre-populated.
    """
    user_id = current_user["user_id"]
    is_admin = current_user.get("role") == "admin"
    db = await _get_db()
    try:
        review = await db.fetchrow(
            "SELECT id, status, user_id, filename, contract_type FROM contract_reviews WHERE id = $1", contract_id
        )
        if not review:
            raise HTTPException(404, "Contract not found")
        if review["user_id"] != user_id and not is_admin:
            raise HTTPException(403, "无权操作此审查")

        if review["status"] not in ("paused_waiting", "completed"):
            logger.warning(
                f"Human decision submitted for review {contract_id} "
                f"in status {review['status']} (expected paused_waiting)"
            )

        # Save decisions to DB
        decision_count = 0
        for d in req.decisions:
            await db.execute(
                """
                INSERT INTO human_review_decisions (review_id, clause_id, action, modified_level, modified_score, comment, skip_revision, decided_by)
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
                """,
                contract_id,
                d.get("clause_id", d.get("rule_id", "")),
                d.get("action", "approve"),
                d.get("modified_level"),
                d.get("modified_score"),
                d.get("comment", ""),
                d.get("skip_revision", False),
                user_id,
            )
            decision_count += 1

        # Update review status to continue
        await db.execute(
            "UPDATE contract_reviews SET status = 'revising', "
            "needs_human_review = FALSE, human_review_status = 'completed', "
            "updated_at = NOW() WHERE id = $1",
            contract_id,
        )

        logger.info(
            f"Human decisions saved for contract {contract_id}: "
            f"{decision_count} decisions"
        )

        # Trigger pipeline resume (continue from revise node)
        # Load the saved state and resume
        from backend.agents.contract_review.graph import run_contract_review_direct

        # Load review info
        saved_path = UPLOAD_DIR / review.get("filename", "")
        text = ""
        if saved_path.exists():
            ext = Path(review.get("filename", "")).suffix.lower()
            if ext == ".txt":
                text = saved_path.read_text(encoding="utf-8", errors="replace")
            elif ext == ".pdf":
                try:
                    import pdfplumber
                    with pdfplumber.open(str(saved_path)) as pdf:
                        parts = [p.extract_text() or "" for p in pdf.pages]
                        text = "\n\n".join(parts)
                except ImportError:
                    pass
            elif ext == ".docx":
                try:
                    from docx import Document
                    doc = Document(str(saved_path))
                    parts = [p.text for p in doc.paragraphs if p.text.strip()]
                    text = "\n\n".join(parts)
                except ImportError:
                    pass

        # Load existing state from DB
        clauses_rows = await db.fetch(
            "SELECT * FROM contract_clauses WHERE review_id = $1 ORDER BY seq_no",
            contract_id,
        )
        cards_rows = await db.fetch(
            "SELECT * FROM contract_review_cards WHERE review_id = $1",
            contract_id,
        )
        evidence_rows = await db.fetch(
            "SELECT * FROM contract_evidence WHERE review_id = $1",
            contract_id,
        )
        rule_rows = await db.fetch(
            "SELECT * FROM rule_findings WHERE review_id = $1",
            contract_id,
        )

        # Run revise node + complete
        from backend.agents.contract_review.revision_writer import revision_writer_node

        state = {
            "contract_id": contract_id,
            "user_id": user_id,
            "text": text,
            "contract_type": review.get("contract_type", ""),
            "clauses": [dict(r) for r in clauses_rows],
            "rule_findings": [dict(r) for r in rule_rows],
            "review_cards": [dict(r) for r in cards_rows],
            "evidence_map": {},
            "human_review_decisions": req.decisions,
            "human_review_status": "completed",
        }
        # Build evidence_map from evidence rows
        for e_row in evidence_rows:
            e = dict(e_row)
            cid = e.get("clause_id", "")
            if cid not in state["evidence_map"]:
                state["evidence_map"][cid] = []
            state["evidence_map"][cid].append(e)

        try:
            r4 = await revision_writer_node(state)  # type: ignore[arg-type]
            from backend.agents.contract_review.graph import _persist_revisions
            if r4.get("revisions"):
                await _persist_revisions(contract_id, r4["revisions"])

            await db.execute(
                "UPDATE contract_reviews SET status = 'completed', updated_at = NOW() WHERE id = $1",
                contract_id,
            )
        except Exception as e:
            logger.exception(f"Resume pipeline failed for contract {contract_id}")
            await db.execute(
                "UPDATE contract_reviews SET status = 'failed', error_message = $2, updated_at = NOW() WHERE id = $1",
                contract_id, str(e)[:500],
            )

        return {
            "status": "resumed",
            "contract_id": contract_id,
            "decisions_saved": decision_count,
        }
    finally:
        await db.close()
