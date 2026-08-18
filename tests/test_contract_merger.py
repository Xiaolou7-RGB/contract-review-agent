"""Tests for contract_merger — 合并已采纳修订，生成修订后合同。"""
from backend.core.contract_merger import (
    merge_clauses,
    build_contract_text,
    build_contract_docx,
    _is_label_title,
    _extract_contract_title,
)


def _clauses():
    return [
        {"seq_no": 1, "clause_id": "c1", "title": "第一条 价款", "content": "原价 100 元"},
        {"seq_no": 2, "clause_id": "c2", "title": "第二条 违约", "content": "违约金 50%"},
        {"seq_no": 3, "clause_id": "c3", "title": "第三条 其他", "content": "其他条款"},
    ]


def _clauses_with_labels():
    """模拟真实合同：含「合同标题/当事人信息/签署落款」等标签类 clause。"""
    return [
        {"seq_no": 1, "clause_id": "t1", "title": "合同标题", "content": "房屋租赁合同"},
        {"seq_no": 2, "clause_id": "t2", "title": "合同编号", "content": "合同编号：GZ-001"},
        {"seq_no": 3, "clause_id": "t3", "title": "当事人信息", "content": "出租人（甲方）：xxx"},
        {"seq_no": 4, "clause_id": "t4", "title": "前言声明", "content": "甲方系合法所有权人"},
        {"seq_no": 5, "clause_id": "c5", "title": "第一条 租赁房屋", "content": "1.1 xxx"},
        {"seq_no": 6, "clause_id": "t6", "title": "签署落款", "content": "甲方（盖章）：xxx"},
    ]


def test_partial_accept():
    merged = merge_clauses(_clauses(), [{"clause_id": "c2", "after_text": "违约金 20%"}])
    assert [m["content"] for m in merged] == ["原价 100 元", "违约金 20%", "其他条款"]
    assert [m["changed"] for m in merged] == [False, True, False]


def test_all_accept():
    revs = [
        {"clause_id": "c1", "after_text": "原价 80"},
        {"clause_id": "c2", "after_text": "违约金 20%"},
        {"clause_id": "c3", "after_text": "其他改"},
    ]
    merged = merge_clauses(_clauses(), revs)
    assert all(m["changed"] for m in merged)


def test_no_accept_keeps_original():
    merged = merge_clauses(_clauses(), [])
    assert all(not m["changed"] for m in merged)
    assert [m["content"] for m in merged] == ["原价 100 元", "违约金 50%", "其他条款"]


def test_rejected_or_needs_lawyer_ignored():
    # 调用方只传 accepted；这里验证即便传入空列表也不会误替换
    merged = merge_clauses(_clauses(), [])
    assert all(not m["changed"] for m in merged)


def test_unmatched_clause_id_keeps_original():
    # 老数据 clause_id 对不上时，不替换、不报错
    merged = merge_clauses(_clauses(), [{"clause_id": "旧格式", "after_text": "不应生效"}])
    assert all(not m["changed"] for m in merged)


def test_build_text_contains_replaced_content():
    merged = merge_clauses(_clauses(), [{"clause_id": "c2", "after_text": "违约金 20%"}])
    text = build_contract_text(merged, "测试合同")
    assert "测试合同" in text
    assert "违约金 20%" in text
    assert "违约金 50%" not in text


def test_build_docx_is_valid_zip():
    merged = merge_clauses(_clauses(), [])
    b = build_contract_docx(merged, "测试合同")
    assert b[:2] == b"PK"  # docx 本质是 zip
    assert len(b) > 1000


# ── 标签过滤与标题提取 ─────────────────────────────────────

def test_is_label_title():
    assert _is_label_title("合同标题") is True
    assert _is_label_title("当事人信息") is True
    assert _is_label_title("前言声明") is True
    assert _is_label_title("签署落款") is True
    assert _is_label_title("第一条 租赁期限") is False
    assert _is_label_title("第二条 违约责任") is False


def test_extract_contract_title():
    assert _extract_contract_title(_clauses_with_labels()) == "房屋租赁合同"


def test_build_text_filters_labels():
    merged = merge_clauses(_clauses_with_labels(), [])
    text = build_contract_text(merged)
    assert "房屋租赁合同" in text          # 真实标题唯一出现
    assert "合同标题" not in text          # 标签被过滤
    assert "当事人信息" not in text
    assert "签署落款" not in text
    assert "第一条 租赁房屋" in text       # 真实条款标题保留
    assert "出租人（甲方）：xxx" in text   # 内容保留


def test_build_docx_filters_labels():
    merged = merge_clauses(_clauses_with_labels(), [])
    b = build_contract_docx(merged)
    from docx import Document
    import io
    doc = Document(io.BytesIO(b))
    texts = [p.text.strip() for p in doc.paragraphs]
    assert "合同标题" not in texts
    assert "当事人信息" not in texts
    assert "签署落款" not in texts
    assert "房屋租赁合同" in texts
    assert "第一条 租赁房屋" in texts
    # 有页眉页脚（header/footer 有内容）
    assert doc.sections[0].header.paragraphs[0].text == "房屋租赁合同"
